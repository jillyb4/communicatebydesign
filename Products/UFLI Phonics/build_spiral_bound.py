#!/usr/bin/env python3
"""
build_spiral_bound.py
=====================
Builds the UFLI Foundations Complete Spiral-Bound Book.

Combines Teacher Guide + Communication Partner Guide + Lessons 5-34 into
a single document optimized for coil binding at 9 Cent Copy.

Structure:
  - Color cover (separate file for color printing)
  - Table of Contents
  - SECTION A: Teacher Guide (imported from existing .docx)
  - SECTION B: Communication Partner Guide (imported from existing .docx)
  - SECTION C: Lessons 5-34 (text-only word reference tables, NO symbol images)
  - Back matter

Print spec (9 Cent Copy — Coil Bound Booklets):
  - Page size: 8.5" x 11" (US Letter)
  - Binding: Coil/spiral on left edge
  - Left margin: 0.75" (0.5" for coil punch + 0.25" inner margin)
  - Right margin: 0.5"
  - Top/bottom: 0.5"
  - Interior: B&W
  - Cover: Color (separate file)
  - Double-sided printing

Dependencies: pip install python-docx --break-system-packages

Usage: python3 build_spiral_bound.py
"""

import json
import os
import sys
import copy
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ── Paths ──
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
OUTPUT_DIR = os.path.join(SCRIPT_DIR, 'Output')
TEACHER_GUIDE = os.path.join(OUTPUT_DIR,
    'UFLI_Teacher_Guide_and_Communication_Partner_Guide.docx')
LESSONS_JSON = '/sessions/funny-lucid-bardeen/ufli_lessons_all.json'

# ── Brand ──
NAVY = RGBColor(0x1B, 0x1F, 0x3B)
TEAL = RGBColor(0x00, 0x6D, 0xA0)
AMBER = RGBColor(0xFF, 0xB7, 0x03)
BLACK = RGBColor(0x00, 0x00, 0x00)
GRAY = RGBColor(0x66, 0x66, 0x66)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

# Fitzgerald Key colors (for text labels in B&W-friendly format)
FITZ_CATS = {
    'People': 'People / Pronouns',
    'Actions': 'Verbs / Actions',
    'Descriptions': 'Descriptions',
    'Nouns': 'Nouns',
    'Prepositions': 'Little Words',
    'Social': 'Social / Feelings',
}

# Fitzgerald classification sets (simplified)
FITZ_PEOPLE = {
    'i', 'you', 'he', 'she', 'we', 'they', 'my', 'your', 'our', 'their',
    'who', 'him', 'his', 'us', 'it', 'me', 'that', 'her', 'them', 'its',
}
FITZ_ACTIONS = {
    'think', 'feel', 'know', 'want', 'need', 'help', 'stop', 'fight', 'change',
    'show', 'make', 'do', 'go', 'get', 'give', 'like', 'live', 'care', 'move',
    'read', 'write', 'work', 'use', 'find', 'try', 'put', 'take', 'see', 'look',
    'hear', 'play', 'come', 'turn', 'open', 'close', 'start', 'finish', 'keep',
    'bring', 'eat', 'drink', 'sit', 'stand', 'walk', 'run', 'jump', 'wait',
    'watch', 'listen', 'talk', 'ask', 'tell', 'say', 'let', 'set', 'fit', 'hit',
    'bit', 'got', 'had', 'has', 'can', 'is', 'am', 'was', 'did', 'hid',
    'nap', 'tap', 'pat', 'sip', 'pin', 'tip', 'nip', 'fan', 'dig', 'dip',
    'dab', 'jog', 'jab', 'beg', 'bet', 'bid', 'bop', 'bug', 'cab', 'cap',
    'cop', 'cod', 'con', 'cot', 'cub', 'cup', 'cut', 'den', 'dog', 'dot',
    'dub', 'dug', 'fad', 'fig', 'fin', 'fix', 'fox', 'fun', 'gab', 'gap',
    'gas', 'gig', 'gum', 'gut', 'hip', 'hog', 'hop', 'hot', 'hug', 'hum',
    'hut', 'jag', 'jam', 'jet', 'jig', 'job', 'jot', 'jug', 'jut', 'kid',
    'kin', 'kit', 'lag', 'lap', 'led', 'leg', 'lid', 'lip', 'lit', 'log',
    'lot', 'lug', 'mad', 'man', 'map', 'mat', 'met', 'mix', 'mob', 'mod',
    'mop', 'mud', 'mug', 'nab', 'nag', 'net', 'nod', 'not', 'nut', 'pad',
    'pan', 'peg', 'pen', 'pet', 'pig', 'pin', 'pit', 'pop', 'pot', 'pug',
    'pun', 'rag', 'rap', 'rat', 'red', 'rid', 'rig', 'rim', 'rip', 'rod',
    'rug', 'run', 'rut', 'sad', 'sag', 'sap', 'sat', 'set', 'sip', 'six',
    'sob', 'sod', 'sub', 'sum', 'tab', 'tad', 'tag', 'tan', 'tax', 'ten',
    'tin', 'top', 'tot', 'tub', 'tug', 'van', 'vat', 'vet', 'wag', 'wax',
    'web', 'wed', 'wet', 'wig', 'win', 'wit', 'yam', 'yet', 'yep', 'zap', 'zip',
}
FITZ_DESCRIPTIONS = {
    'good', 'bad', 'right', 'wrong', 'different', 'same', 'more', 'less',
    'big', 'small', 'little', 'new', 'old', 'hard', 'easy', 'long', 'short',
    'fast', 'slow', 'happy', 'sad', 'hot', 'cold', 'all', 'some', 'not',
    'other', 'next', 'last', 'first', 'many', 'few',
}
FITZ_PREPOSITIONS = {
    'because', 'before', 'after', 'then', 'but', 'if', 'about', 'and',
    'when', 'where', 'in', 'on', 'at', 'up', 'down', 'to', 'from', 'with',
    'for', 'of', 'by', 'out', 'off', 'over', 'under', 'or', 'so', 'just',
    'very', 'again', 'still', 'also', 'too', 'yet',
}
FITZ_SOCIAL = {
    'yes', 'no', 'please', 'thank', 'sorry', 'hi', 'hello', 'bye',
    'why', 'what', 'how', 'okay', 'wow', 'oh',
}


def classify(word):
    w = word.lower().strip()
    if w in FITZ_PEOPLE: return 'People'
    if w in FITZ_ACTIONS: return 'Actions'
    if w in FITZ_DESCRIPTIONS: return 'Descriptions'
    if w in FITZ_PREPOSITIONS: return 'Prepositions'
    if w in FITZ_SOCIAL: return 'Social'
    return 'Nouns'


# ══════════════════════════════════════════════════════════════
# Document helpers
# ══════════════════════════════════════════════════════════════

def set_cell_shading(cell, hex_color):
    """Set cell background color (removes existing shading first)."""
    tcPr = cell._tc.get_or_add_tcPr()
    existing = tcPr.findall(qn('w:shd'))
    for e in existing:
        tcPr.remove(e)
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    tcPr.append(shading)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = NAVY
        run.font.name = 'Arial'
    return h


def add_para(doc, text, bold=False, size=11, color=BLACK, space_after=Pt(4)):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    p.paragraph_format.space_after = space_after
    return p


def add_page_break(doc):
    doc.add_page_break()


# ══════════════════════════════════════════════════════════════
# Import Teacher Guide content
# ══════════════════════════════════════════════════════════════

def import_teacher_guide(target_doc):
    """Copy all content from the Teacher Guide .docx into target doc."""
    if not os.path.exists(TEACHER_GUIDE):
        print(f'  ⚠️  Teacher Guide not found: {TEACHER_GUIDE}')
        return

    source = Document(TEACHER_GUIDE)
    print(f'  Importing Teacher Guide: {len(source.paragraphs)} paragraphs, {len(source.tables)} tables')

    for element in source.element.body:
        # Deep copy each element from source to target
        new_element = copy.deepcopy(element)
        target_doc.element.body.append(new_element)

    print('  ✅ Teacher Guide imported')


# ══════════════════════════════════════════════════════════════
# Session Tracker — landscape partner handoff page
# ══════════════════════════════════════════════════════════════

def _remove_cell_borders(cell):
    """Remove all borders from a table cell (removes existing tcBorders first)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    existing = tcPr.findall(qn('w:tcBorders'))
    for e in existing:
        tcPr.remove(e)
    tcBorders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        '<w:top w:val="none" w:sz="0" w:space="0"/>'
        '<w:left w:val="none" w:sz="0" w:space="0"/>'
        '<w:bottom w:val="none" w:sz="0" w:space="0"/>'
        '<w:right w:val="none" w:sz="0" w:space="0"/>'
        '</w:tcBorders>')
    tcPr.append(tcBorders)


def _set_col_width(table, col_idx, width):
    """Set width for all cells in a column."""
    for row in table.rows:
        row.cells[col_idx].width = width


def _force_table_widths(table, col_widths_inches):
    """Force exact column widths via tblGrid XML and disable autofit.

    python-docx cell.width only sets 'preferred' widths. Word may override.
    This sets the tblGrid directly, which Word respects.
    """
    tbl = table._tbl
    # Disable autofit
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    if tbl.tblPr is None:
        tbl.insert(0, tblPr)
    # Remove existing layout
    for existing in tblPr.findall(qn('w:tblLayout')):
        tblPr.remove(existing)
    tblPr.append(parse_xml(f'<w:tblLayout {nsdecls("w")} w:type="fixed"/>'))

    # Set total table width
    total_emu = sum(int(w * 914400) for w in col_widths_inches)
    total_dxa = int(total_emu / 635)
    for existing in tblPr.findall(qn('w:tblW')):
        tblPr.remove(existing)
    tblPr.append(parse_xml(f'<w:tblW {nsdecls("w")} w:w="{total_dxa}" w:type="dxa"/>'))

    # Replace tblGrid
    for existing in tbl.findall(qn('w:tblGrid')):
        tbl.remove(existing)
    grid_xml = f'<w:tblGrid {nsdecls("w")}>'
    for w in col_widths_inches:
        dxa = int(w * 914400 / 635)
        grid_xml += f'<w:gridCol w:w="{dxa}"/>'
    grid_xml += '</w:tblGrid>'
    tbl.insert(1, parse_xml(grid_xml))  # after tblPr

    # Also set each cell width for consistency
    for ci, w in enumerate(col_widths_inches):
        dxa = int(w * 914400 / 635)
        for row in table.rows:
            tc = row.cells[ci]._tc
            tcPr = tc.get_or_add_tcPr()
            for existing in tcPr.findall(qn('w:tcW')):
                tcPr.remove(existing)
            tcPr.append(parse_xml(f'<w:tcW {nsdecls("w")} w:w="{dxa}" w:type="dxa"/>'))
            tc.width = Emu(int(w * 914400))


def _set_row_height(row, height_pt):
    """Set exact row height (removes existing trHeight first)."""
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    existing = trPr.findall(qn('w:trHeight'))
    for e in existing:
        trPr.remove(e)
    trHeight = parse_xml(
        f'<w:trHeight {nsdecls("w")} w:val="{int(height_pt * 20)}" w:hRule="exact"/>')
    trPr.append(trHeight)


def _set_cell_vertical_alignment(cell, align='center'):
    """Set vertical alignment of cell content (removes existing vAlign first)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    existing = tcPr.findall(qn('w:vAlign'))
    for e in existing:
        tcPr.remove(e)
    vAlign = parse_xml(f'<w:vAlign {nsdecls("w")} w:val="{align}"/>')
    tcPr.append(vAlign)


def _compact_cell(cell, text, size=8, bold=False, color=BLACK, align=None):
    """Write text into a cell with zero-padding paragraph formatting."""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    if align:
        p.alignment = align
    r = p.add_run(text)
    r.font.name = 'Arial'
    r.font.size = Pt(size)
    r.bold = bold
    r.font.color.rgb = color


def build_tracker_page(doc, lesson_num, lesson_data=None):
    """One landscape page per lesson — the partner handoff tool.

    Fills the entire 11×8.5 landscape page. This is the working document
    that travels with the student. Every partner picks it up and knows
    exactly what happened last session and what to do next.

    Available width:  11 - 0.75 (spiral) - 0.4 = 9.85"
    Available height: 8.5 - 0.35 - 0.35 = 7.8"
    """
    PAGE_CONTENT_W = Inches(9.85)  # usable width

    phoneme = ''
    grapheme = ''
    new_words = []
    heart_words = []
    if lesson_data:
        phoneme = lesson_data.get('phoneme', '')
        grapheme = lesson_data.get('grapheme', '')
        new_words = lesson_data.get('newWords', [])
        heart_words = lesson_data.get('heartWords', [])

    # Force landscape
    new_section = doc.add_section(start_type=2)
    new_section.orientation = WD_ORIENT.LANDSCAPE
    new_section.page_width = Inches(11)
    new_section.page_height = Inches(8.5)
    new_section.left_margin = Inches(0.75)
    new_section.right_margin = Inches(0.4)
    new_section.top_margin = Inches(0.35)
    new_section.bottom_margin = Inches(0.35)

    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    # HEADER — lesson title + session tracker label
    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    lesson_label = f'Lesson {lesson_num}'
    if phoneme:
        lesson_label += f'  —  {phoneme}  ({grapheme})'

    h = doc.add_paragraph()
    h.paragraph_format.space_after = Pt(2)
    h.paragraph_format.space_before = Pt(0)
    hr = h.add_run(lesson_label)
    hr.font.name = 'Arial'; hr.font.size = Pt(16); hr.font.color.rgb = NAVY; hr.bold = True
    hr2 = h.add_run('     Session Tracker')
    hr2.font.name = 'Arial'; hr2.font.size = Pt(11); hr2.font.color.rgb = TEAL; hr2.bold = True

    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    # INFO BAR — student/partner/role/date as a borderless table
    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    info_t = doc.add_table(rows=1, cols=4)
    info_t.style = 'Table Grid'
    info_t.alignment = WD_TABLE_ALIGNMENT.LEFT
    info_labels = ['Student:', 'Partner:', 'Role:', 'Date:']
    info_widths = [3.0, 3.0, 2.0, 1.85]
    _force_table_widths(info_t, info_widths)
    for ci, lbl_text in enumerate(info_labels):
        cell = info_t.rows[0].cells[ci]
        _compact_cell(cell, lbl_text, size=9, bold=True, color=NAVY)
        _set_cell_vertical_alignment(cell, 'center')
    _set_row_height(info_t.rows[0], 22)

    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    # FOCUS BAR — tools, access, steps, pacing as a 2-row table
    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    focus_t = doc.add_table(rows=2, cols=4)
    focus_t.style = 'Table Grid'
    focus_t.alignment = WD_TABLE_ALIGNMENT.LEFT
    focus_widths = [4.0, 2.6, 1.5, 1.75]
    _force_table_widths(focus_t, focus_widths)

    # Row 1: labels (navy background)
    focus_headers = ['Tools (circle)', 'Access (circle)', 'Steps (circle)', 'Pacing (circle)']
    for ci, fh in enumerate(focus_headers):
        cell = focus_t.rows[0].cells[ci]
        _compact_cell(cell, fh, size=7, bold=True, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(cell, '1B1F3B')
        _set_cell_vertical_alignment(cell, 'center')
    _set_row_height(focus_t.rows[0], 16)

    # Row 2: options to circle
    focus_opts = [
        'Alt Pencil    E-Trans    Symbol Cards    Letter Cards    SGD',
        'Direct    Eye Gaze    Partner Scan    Switch',
        '1    2    3    4    5    6    7    8',
        'Full    Partial (to Step ___)\nReview    Reteach',
    ]
    for ci, fo in enumerate(focus_opts):
        cell = focus_t.rows[1].cells[ci]
        _compact_cell(cell, fo, size=8, align=WD_ALIGN_PARAGRAPH.CENTER)
        _set_cell_vertical_alignment(cell, 'center')
    _set_row_height(focus_t.rows[1], 26)

    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    # PROMPT LEVEL TABLE — full width, generous row heights
    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    lbl_p = doc.add_paragraph()
    lbl_p.paragraph_format.space_before = Pt(6)
    lbl_p.paragraph_format.space_after = Pt(2)
    lr = lbl_p.add_run('Prompt Level  ')
    lr.font.name = 'Arial'; lr.font.size = Pt(10); lr.bold = True; lr.font.color.rgb = NAVY
    lr2 = lbl_p.add_run('(circle one per row)    I = Independent    3 = CTD 3s    5 = CTD 5s    M = Full Model    — = Not presented')
    lr2.font.name = 'Arial'; lr2.font.size = Pt(8); lr2.font.color.rgb = GRAY

    steps = ['1. Phonemic Awareness', '2. Visual Drill', '3. Auditory Drill',
             '4. Blending Drill', '5. New Concept', '6. Word Work',
             '7. Heart Words', '8. Connected Text']
    levels = ['I', '3', '5', 'M', '—']

    pt = doc.add_table(rows=len(steps) + 1, cols=7)
    pt.style = 'Table Grid'
    pt.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Full-width columns: step 2.5", levels 0.8" each, notes fills remainder
    pt_widths = [2.5] + [0.8] * 5 + [3.35]
    _force_table_widths(pt, pt_widths)

    # Header row
    for i, h_text in enumerate(['UFLI Step'] + levels + ['Notes']):
        cell = pt.rows[0].cells[i]
        _compact_cell(cell, h_text, size=8, bold=True, color=WHITE,
                      align=WD_ALIGN_PARAGRAPH.CENTER if i > 0 else None)
        set_cell_shading(cell, '1B1F3B')
        _set_cell_vertical_alignment(cell, 'center')
    _set_row_height(pt.rows[0], 18)

    # Data rows — generous height for circling
    for ri, step in enumerate(steps):
        row = pt.rows[ri + 1]
        shade = 'F0F2F4' if ri % 2 == 0 else 'FFFFFF'
        _set_row_height(row, 26)

        _compact_cell(row.cells[0], step, size=9, bold=True, color=NAVY)
        set_cell_shading(row.cells[0], shade)
        _set_cell_vertical_alignment(row.cells[0], 'center')

        for li, lv in enumerate(levels):
            _compact_cell(row.cells[li + 1], lv, size=12, color=NAVY,
                          align=WD_ALIGN_PARAGRAPH.CENTER)
            set_cell_shading(row.cells[li + 1], shade)
            _set_cell_vertical_alignment(row.cells[li + 1], 'center')

        set_cell_shading(row.cells[6], shade)

    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    # CORE WORD USE — full width, pre-populated
    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    core_words = [w['word'] for w in new_words if w.get('type') == 'core'] if new_words else []

    cwl = doc.add_paragraph()
    cwl.paragraph_format.space_before = Pt(6)
    cwl.paragraph_format.space_after = Pt(2)
    cwr = cwl.add_run('Core Word Use  ')
    cwr.font.name = 'Arial'; cwr.font.size = Pt(10); cwr.bold = True; cwr.font.color.rgb = NAVY
    cwr2 = cwl.add_run('✓ spontaneous    M modeled    — not observed    ★ generalized')
    cwr2.font.name = 'Arial'; cwr2.font.size = Pt(8); cwr2.font.color.rgb = GRAY

    # Always 8 cols to fill the width, pad or truncate
    n_cols = 8
    display_words = (core_words + [''] * n_cols)[:n_cols]
    cw_each = 9.85 / n_cols

    cw_t = doc.add_table(rows=2, cols=n_cols)
    cw_t.style = 'Table Grid'
    cw_t.alignment = WD_TABLE_ALIGNMENT.LEFT
    _force_table_widths(cw_t, [cw_each] * n_cols)

    for ci, w in enumerate(display_words):
        _compact_cell(cw_t.rows[0].cells[ci], f'★ {w}' if w else '', size=9, bold=True, color=NAVY,
                      align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(cw_t.rows[0].cells[ci], 'F0F2F4')
        _set_cell_vertical_alignment(cw_t.rows[0].cells[ci], 'center')
        _set_cell_vertical_alignment(cw_t.rows[1].cells[ci], 'center')
    _set_row_height(cw_t.rows[0], 22)
    _set_row_height(cw_t.rows[1], 22)

    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    # READING & CONNECTED TEXT + HEART WORDS
    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    rdg = doc.add_paragraph()
    rdg.paragraph_format.space_before = Pt(6)
    rdg.paragraph_format.space_after = Pt(1)
    rr = rdg.add_run('Reading & Connected Text   ')
    rr.font.name = 'Arial'; rr.font.size = Pt(10); rr.bold = True; rr.font.color.rgb = NAVY
    rt = rdg.add_run(
        '☐ Attempted passage    ☐ Not attempted    '
        'Accuracy: ___ / ___    '
        '☐ Independent    ☐ Partner read-along    ☐ Symbol support    '
        '☐ Self-corrected    ☐ Tracked L→R    ☐ Re-read'
    )
    rt.font.name = 'Arial'; rt.font.size = Pt(8)

    if heart_words:
        hw_p = doc.add_paragraph()
        hw_p.paragraph_format.space_before = Pt(1)
        hw_p.paragraph_format.space_after = Pt(1)
        hwr = hw_p.add_run('Heart Words This Lesson:   ')
        hwr.font.name = 'Arial'; hwr.font.size = Pt(9); hwr.bold = True; hwr.font.color.rgb = NAVY
        hw_names = [hw if isinstance(hw, str) else hw.get('word', str(hw)) for hw in heart_words]
        hwr2 = hw_p.add_run('     '.join(f'♥ {w}' for w in hw_names))
        hwr2.font.name = 'Arial'; hwr2.font.size = Pt(9)

    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    # NOTES FOR NEXT COMMUNICATION PARTNER — generous space
    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    np_p = doc.add_paragraph()
    np_p.paragraph_format.space_before = Pt(6)
    np_p.paragraph_format.space_after = Pt(2)
    nr = np_p.add_run('Notes for Next Communication Partner   ')
    nr.font.name = 'Arial'; nr.font.size = Pt(10); nr.bold = True; nr.font.color.rgb = NAVY
    nr2 = np_p.add_run('What went well?  What was hard?  What should the next person know?')
    nr2.font.name = 'Arial'; nr2.font.size = Pt(8); nr2.font.color.rgb = GRAY

    for _ in range(4):
        lp = doc.add_paragraph()
        lp.paragraph_format.space_before = Pt(6)
        lp.paragraph_format.space_after = Pt(0)
        lr = lp.add_run('_' * 150)
        lr.font.name = 'Arial'; lr.font.size = Pt(7); lr.font.color.rgb = RGBColor(0xBB, 0xBB, 0xBB)

    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    # MASTERY DECISION — anchored at bottom
    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    mp = doc.add_paragraph()
    mp.paragraph_format.space_before = Pt(8)
    mp.paragraph_format.space_after = Pt(0)
    mr = mp.add_run('Mastery Decision:   ☐ Move to next lesson     ☐ Reteach this lesson     ☐ Adjust access method     ☐ Consult with team')
    mr.font.name = 'Arial'; mr.font.size = Pt(10); mr.bold = True

    # ── Switch BACK to portrait ──
    next_section = doc.add_section(start_type=2)
    next_section.orientation = WD_ORIENT.PORTRAIT
    next_section.page_width = Inches(8.5)
    next_section.page_height = Inches(11)
    next_section.left_margin = Inches(0.75)
    next_section.right_margin = Inches(0.5)
    next_section.top_margin = Inches(0.5)
    next_section.bottom_margin = Inches(0.5)


# ══════════════════════════════════════════════════════════════
# Build lesson pages (text only — no symbol images)
# ══════════════════════════════════════════════════════════════

def _keep_table_together(table):
    """Prevent a table from splitting across pages."""
    for row in table.rows:
        tr = row._tr
        trPr = tr.get_or_add_trPr()
        # cantSplit = keep row on one page
        for existing in trPr.findall(qn('w:cantSplit')):
            trPr.remove(existing)
        trPr.append(parse_xml(f'<w:cantSplit {nsdecls("w")}/>'))


def _add_step_row(table, step_label, time, what_to_do, shaded=False):
    """Add one row to the Run This Lesson table."""
    row = table.add_row()
    shade = 'F0F4F8' if shaded else 'FFFFFF'

    # Step cell
    cell0 = row.cells[0]
    cell0.text = ''
    p0 = cell0.paragraphs[0]
    r0 = p0.add_run(step_label)
    r0.font.name = 'Arial'; r0.font.size = Pt(8); r0.bold = True; r0.font.color.rgb = NAVY
    p0.paragraph_format.space_after = Pt(0)
    tp = cell0.add_paragraph()
    tr = tp.add_run(time)
    tr.font.name = 'Arial'; tr.font.size = Pt(7); tr.font.color.rgb = GRAY
    tp.paragraph_format.space_after = Pt(0)
    set_cell_shading(cell0, shade)
    _set_cell_vertical_alignment(cell0, 'top')

    # What to do cell
    cell1 = row.cells[1]
    cell1.text = ''
    p1 = cell1.paragraphs[0]
    r1 = p1.add_run(what_to_do)
    r1.font.name = 'Arial'; r1.font.size = Pt(8)
    p1.paragraph_format.space_after = Pt(0)
    set_cell_shading(cell1, shade)
    _set_cell_vertical_alignment(cell1, 'top')


def build_lesson_section(doc, lesson):
    """Build one lesson's pages for the spiral book.

    Each lesson = 3 pages:
      Page 1 (portrait): RUN THIS LESSON — 8 steps with this lesson's actual words
      Page 2 (portrait): Word Reference — word tables, binder pull list, heart words
      Page 3 (landscape): Session Tracker — fill in during/after the session
    """
    num = lesson['number']
    phoneme = lesson.get('phoneme', '')
    grapheme = lesson.get('grapheme', '')
    new_words = lesson.get('newWords', [])
    review_words = lesson.get('reviewWords', [])
    heart_words = lesson.get('heartWords', [])
    morph_notes = lesson.get('morphologyNotes', [])

    is_letter_only = not new_words and not review_words  # Lessons 1-4

    # ══════════════════════════════════════════════════════════
    # PAGE 1: RUN THIS LESSON (portrait)
    # Everything the partner needs to run the lesson — step by step.
    # Lessons 1-4: letter introduction routine
    # Lessons 5+: full 8-step routine with this lesson's actual words
    # ══════════════════════════════════════════════════════════

    add_page_break(doc)

    # ── Lesson header with phoneme ──
    title = f'Lesson {num}'
    if phoneme:
        title += f'  —  {phoneme}  ({grapheme})'
    add_heading(doc, title, level=1)

    if is_letter_only:
        # ── LESSONS 1-4: LETTER INTRODUCTION ──
        add_para(doc, 'Letter Introduction — Run This Lesson',
                 bold=True, size=12, color=TEAL, space_after=Pt(4))
        add_para(doc,
            f'This lesson introduces the letter {grapheme}. Words begin in Lesson 5. '
            f'Today builds letter recognition and the alternative pencil routine. '
            f'Your student starts with the class — they do not wait.',
            size=9, color=GRAY, space_after=Pt(6))

        # Setup checklist
        setup = doc.add_table(rows=1, cols=2)
        setup.style = 'Table Grid'
        setup.alignment = WD_TABLE_ALIGNMENT.LEFT
        _force_table_widths(setup, [3.5, 3.75])

        left_c = setup.rows[0].cells[0]
        left_c.text = ''
        lp = left_c.paragraphs[0]
        lr = lp.add_run('Have Ready:')
        lr.font.name = 'Arial'; lr.font.size = Pt(9); lr.bold = True; lr.font.color.rgb = NAVY
        for t in [
            f'☐ Letter card: {grapheme.upper()} / {grapheme.lower()}',
            f'☐ Phoneme card: {phoneme}',
            '☐ Alternative pencil (progressive reveal)',
            '☐ E-trans board',
            '☐ Sound access (partner voice or device)',
        ]:
            tp = left_c.add_paragraph()
            tr = tp.add_run(t)
            tr.font.name = 'Arial'; tr.font.size = Pt(8)
            tp.paragraph_format.space_after = Pt(1)
        set_cell_shading(left_c, 'F4F6F8')

        right_c = setup.rows[0].cells[1]
        right_c.text = ''
        rp = right_c.paragraphs[0]
        rr = rp.add_run('Goals for This Lesson:')
        rr.font.name = 'Arial'; rr.font.size = Pt(9); rr.bold = True; rr.font.color.rgb = NAVY
        for g in [
            f'Student recognizes {grapheme.upper()} and {grapheme.lower()}',
            f'Student finds {grapheme} on their alternative pencil',
            f'Student hears {phoneme} and connects it to {grapheme}',
            'Student practices selecting the letter on e-trans board',
            'Build motor plan for letter location on keyboard/board',
        ]:
            gp = right_c.add_paragraph()
            gr = gp.add_run(g)
            gr.font.name = 'Arial'; gr.font.size = Pt(8)
            gp.paragraph_format.space_after = Pt(1)
        set_cell_shading(right_c, 'F4F6F8')

        doc.add_paragraph().paragraph_format.space_after = Pt(4)

        # Step-by-step activities for letter-only lessons
        add_para(doc, 'Step by Step', bold=True, size=11, color=NAVY, space_after=Pt(4))

        step_table = doc.add_table(rows=1, cols=2)
        step_table.style = 'Table Grid'
        step_table.alignment = WD_TABLE_ALIGNMENT.LEFT
        _force_table_widths(step_table, [1.5, 5.75])

        # Header
        for i, h in enumerate(['Step', 'What to Do']):
            cell = step_table.rows[0].cells[i]
            cell.text = h
            set_cell_shading(cell, '1B1F3B')
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.color.rgb = WHITE; r.bold = True
                    r.font.name = 'Arial'; r.font.size = Pt(8)

        _add_step_row(step_table, 'INTRODUCE', '~3 min',
            f'Show the letter card ({grapheme.upper()}/{grapheme.lower()}). '
            f'Say: "This letter is {grapheme}. It makes this sound: {phoneme}." '
            f'Produce {phoneme} clearly. Let student explore: sandpaper letter, '
            f'playdough, tracing, light box. Place phoneme card nearby — '
            f'student presses it to hear {phoneme} while they trace.', True)

        _add_step_row(step_table, 'LETTER MATCHING', '~3 min',
            f'Show {grapheme.upper()}. Student finds {grapheme.lower()} on alternative pencil '
            f'or letter cards. Mix in known letters as distractors. Start with 2 choices, '
            f'build to 3–4. WAIT 5–7 seconds before helping.')

        _add_step_row(step_table, 'SOUND CONNECTION', '~5 min',
            f'Say {phoneme}. Student selects {grapheme} from a field of 2–3 letters '
            f'on the e-trans board. When student selects ANY letter, produce the sound '
            f'that letter makes — right or wrong. This is the auditory loop.', True)

        _add_step_row(step_table, 'PROGRESSIVE REVEAL', '~2 min',
            f'Unhide {grapheme} on the frequency keyboard / alphabet board. '
            f'Student now has access to all letters taught so far. '
            f'Student practices finding {grapheme} on their own system.')

        _add_step_row(step_table, 'ALT PENCIL PRACTICE', '~5 min',
            f'Student uses their access method to select {grapheme} on their '
            f'alternative pencil. Partner confirms by producing {phoneme}. '
            f'Practice 3–5 times. Goal: student builds motor plan for where '
            f'{grapheme} lives on their keyboard/board.', True)

        _keep_table_together(step_table)

        # Auditory loop reminder box
        doc.add_paragraph().paragraph_format.space_after = Pt(4)
        loop_box = doc.add_table(rows=1, cols=1)
        loop_box.alignment = WD_TABLE_ALIGNMENT.LEFT
        _force_table_widths(loop_box, [7.25])
        lbc = loop_box.rows[0].cells[0]
        lbc.text = ''
        lbp = lbc.paragraphs[0]
        lbr = lbp.add_run('AUDITORY LOOP (every interaction): ')
        lbr.font.name = 'Arial'; lbr.font.size = Pt(9); lbr.bold = True; lbr.font.color.rgb = AMBER
        lbr2 = lbp.add_run(
            f'Student selects {grapheme} → you produce {phoneme}. '
            f'Student selects a wrong letter → you produce THAT sound. '
            f'The student hears what they selected to self-monitor. '
            f'Sound, not letter name. Always. Immediately.')
        lbr2.font.name = 'Arial'; lbr2.font.size = Pt(9)
        set_cell_shading(lbc, 'FFF8E7')

        # Tracker page
        build_tracker_page(doc, num, lesson)
        return

    # ══════════════════════════════════════════════════════════
    # LESSONS 5+: PAGE 1 — RUN THIS LESSON
    # The 8 UFLI steps with THIS lesson's actual words.
    # A brand-new para on day 1 reads this page and knows what to do.
    # ══════════════════════════════════════════════════════════

    # Get word lists as readable strings
    new_word_names = [w['word'] for w in new_words]
    review_word_names = [w['word'] for w in review_words] if review_words else []
    hw_names = [hw if isinstance(hw, str) else hw.get('word', str(hw)) for hw in heart_words]
    all_word_names = new_word_names + review_word_names

    # E-trans preview: pick 3-4 most useful words for the partner to visualize
    etrans_words = new_word_names[:4]
    etrans_str = ', '.join(etrans_words)

    add_para(doc, 'Run This Lesson', bold=True, size=12, color=TEAL, space_after=Pt(2))

    # Compact setup line
    setup_line = doc.add_paragraph()
    setup_line.paragraph_format.space_after = Pt(4)
    sr1 = setup_line.add_run('Have ready: ')
    sr1.font.name = 'Arial'; sr1.font.size = Pt(8); sr1.bold = True; sr1.font.color.rgb = NAVY
    sr2 = setup_line.add_run(
        f'symbol cards for {len(new_words)} new words  •  '
        f'{len(review_word_names)} review cards from binder  •  '
        f'alternative pencil  •  e-trans board  •  sound access')
    sr2.font.name = 'Arial'; sr2.font.size = Pt(8); sr2.font.color.rgb = GRAY

    # ── 8-STEP TABLE ──
    step_table = doc.add_table(rows=1, cols=2)
    step_table.style = 'Table Grid'
    step_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    _force_table_widths(step_table, [1.5, 5.75])

    # Header
    for i, h in enumerate(['Step', 'What to Do (with this lesson\'s words)']):
        cell = step_table.rows[0].cells[i]
        cell.text = h
        set_cell_shading(cell, '1B1F3B')
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.color.rgb = WHITE; r.bold = True
                r.font.name = 'Arial'; r.font.size = Pt(8)

    # Step 1: Phonemic Awareness
    _add_step_row(step_table,
        '1. Phonemic\nAwareness', '~2 min',
        f'E-trans board: place 2–4 symbol cards ({etrans_str}). '
        f'Say the phonemes for a word. Student blends internally, then selects the word. '
        f'WAIT 5–7 sec. When student selects: say the word, then segment it back into sounds.',
        True)

    # Step 2: Visual Drill
    _add_step_row(step_table,
        '2. Visual Drill', '~3 min',
        f'Flash grapheme card: {grapheme}. Student indicates the sound {phoneme}. '
        f'Also flash review graphemes. PRODUCE THE SOUND for whatever the student selects — '
        f'right or wrong. Say the sound, not the letter name. Keep it moving.')

    # Step 3: Auditory Drill
    _add_step_row(step_table,
        '3. Auditory Drill', '~5 min',
        f'Say {phoneme}. Student finds {grapheme} on their alternative pencil. '
        f'WAIT 5–7 sec. When student selects a letter: produce the sound that letter makes. '
        f'If correct: confirm. If incorrect: produce the correct sound, point to correct letter.',
        True)

    # Step 4: Blending Drill
    blend_words = ', '.join(new_word_names[:5])
    _add_step_row(step_table,
        '4. Blending Drill', '~5 min',
        f'E-trans board with symbol cards. Show graphemes one at a time. '
        f'"Blend them. What word? Show me." Words: {blend_words}. '
        f'Use phonetically similar words as distractors. WAIT 5–7 sec. '
        f'When student selects: sound it out, then say the whole word.')

    # Step 5: New Concept
    _add_step_row(step_table,
        '5. New Concept', '~15 min',
        f'Today\'s concept: {phoneme} ({grapheme}). '
        f'INTRODUCE: Point to grapheme, produce {phoneme}. Let student explore the letter: '
        f'sandpaper, playdough, tracing. '
        f'DECODE: E-trans with 2–3 cards. "What word? Show me." WAIT 7–10 sec. '
        f'ENCODE: "Spell it on your alphabet." Produce the sound for each letter selected.',
        True)

    # Step 6: Word Work / Word Chain
    chain_words = ', '.join(new_word_names[:3]) if len(new_word_names) >= 3 else ', '.join(new_word_names)
    _add_step_row(step_table,
        '6. Word Work', '~5 min',
        f'Word chain with: {chain_words}. Change one phoneme at a time. '
        f'Produce BOTH sounds clearly when announcing the change. '
        f'Student selects the new word from e-trans or uses letter tiles/pencil. '
        f'WAIT 5–7 sec. If student needs >2 attempts: model, move on, circle back.')

    # Step 7: Heart Words
    if hw_names:
        hw_str = ', '.join(hw_names)
        _add_step_row(step_table,
            '7. Heart Words', '~6 min',
            f'Heart words this lesson: {hw_str}. '
            f'Show the word. Say it, then sound it out letter by letter. '
            f'Identify the regular part and the heart part. '
            f'Student identifies on e-trans, then spells on alternative pencil. '
            f'Practice 3 times during the lesson. WAIT 5 sec per response.',
            True)
    else:
        _add_step_row(step_table,
            '7. Heart Words', '~6 min',
            'No new heart words this lesson. Review previous heart words from the binder. '
            'Flash 3–5 heart word cards. Student identifies each on sight. '
            'If <2 sessions since last heart word: quiz, don\'t re-teach.',
            True)

    # Step 8: Connected Text
    target_str = ', '.join(new_word_names[:6])
    _add_step_row(step_table,
        '8. Connected Text', '~15 min',
        f'BEFORE: Preview target words ({target_str}) with symbol cards on e-trans. '
        f'DURING: Read the decodable passage. PAUSE before target words. WAIT 5–10 sec. '
        f'Student selects the word. Say it naturally (fluency context). '
        f'AFTER: 2–3 comprehension questions using symbol cards. Do NOT skip this step.')

    _keep_table_together(step_table)

    # ── AUDITORY LOOP BOX ──
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    loop_box = doc.add_table(rows=1, cols=1)
    loop_box.alignment = WD_TABLE_ALIGNMENT.LEFT
    _force_table_widths(loop_box, [7.25])
    lbc = loop_box.rows[0].cells[0]
    lbc.text = ''
    lbp = lbc.paragraphs[0]
    lbr = lbp.add_run('AUDITORY LOOP (every step, every interaction): ')
    lbr.font.name = 'Arial'; lbr.font.size = Pt(9); lbr.bold = True; lbr.font.color.rgb = AMBER
    lbr2 = lbp.add_run(
        'Student selects → YOU PRODUCE THE SOUND. Right or wrong. Always. Immediately. '
        'Decoding: say the word, then sound it out. Encoding: say the sound that letter makes. '
        'Connected text: say the word naturally.')
    lbr2.font.name = 'Arial'; lbr2.font.size = Pt(9)
    set_cell_shading(lbc, 'FFF8E7')

    # Morphology note if present
    if morph_notes:
        doc.add_paragraph().paragraph_format.space_after = Pt(2)
        for note in morph_notes:
            if isinstance(note, dict):
                text = f'Morphology teaching moment: {note.get("word", "")} ← {note.get("base", "")}: {note.get("note", "")}'
            else:
                text = f'Morphology teaching moment: {str(note)}'
            mp = doc.add_paragraph()
            mr = mp.add_run(text)
            mr.font.name = 'Arial'; mr.font.size = Pt(8); mr.font.color.rgb = TEAL; mr.italic = True

    # ══════════════════════════════════════════════════════════
    # PAGE 2: WORD REFERENCE (portrait)
    # New words, review words to pull from binder, heart words
    # ══════════════════════════════════════════════════════════
    add_page_break(doc)

    ref_title = f'Lesson {num} — Word Reference'
    add_heading(doc, ref_title, level=2)
    add_para(doc, 'Print the symbol cards from the per-lesson packet. Pull review cards from the binder by tab color.',
             size=9, color=GRAY, space_after=Pt(4))

    # ── NEW WORDS TABLE — organized by Fitzgerald Key category ──
    if new_words:
        add_para(doc, f'New Words ({len(new_words)})', bold=True, size=11, color=TEAL, space_after=Pt(2))

        # Group by Fitzgerald category
        by_cat = {}
        for w in new_words:
            cat = classify(w['word'])
            if cat not in by_cat:
                by_cat[cat] = []
            by_cat[cat].append(w)

        word_table = doc.add_table(rows=1, cols=4)
        word_table.style = 'Table Grid'
        word_table.alignment = WD_TABLE_ALIGNMENT.LEFT
        _force_table_widths(word_table, [1.8, 1.3, 2.2, 1.95])

        # Header
        for i, h in enumerate(['Word', 'Core / Fringe', 'Fitzgerald Category', 'Binder Tab']):
            cell = word_table.rows[0].cells[i]
            cell.text = h
            set_cell_shading(cell, '1B1F3B')
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.color.rgb = WHITE; r.bold = True
                    r.font.name = 'Arial'; r.font.size = Pt(8)

        cat_order = ['People', 'Actions', 'Descriptions', 'Nouns', 'Prepositions', 'Social']
        for cat in cat_order:
            if cat not in by_cat:
                continue
            for w in sorted(by_cat[cat], key=lambda x: x['word'].lower()):
                row = word_table.add_row()
                word_text = w['word']
                if w.get('type') == 'core':
                    word_text = f'★ {word_text}'
                row.cells[0].text = word_text
                row.cells[1].text = w.get('type', 'fringe').capitalize()
                row.cells[2].text = FITZ_CATS.get(cat, cat)
                row.cells[3].text = cat
                for cell in row.cells:
                    for p in cell.paragraphs:
                        for r in p.runs:
                            r.font.name = 'Arial'; r.font.size = Pt(8)

        _keep_table_together(word_table)
        doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # ── REVIEW WORDS — pull from binder, organized by tab ──
    if review_words:
        add_para(doc, f'Review Words — Pull From Binder ({len(review_words)})',
                 bold=True, size=11, color=TEAL, space_after=Pt(2))

        # Group by Fitzgerald category
        by_cat = {}
        for w in review_words:
            cat = classify(w['word'])
            if cat not in by_cat:
                by_cat[cat] = []
            by_cat[cat].append(w['word'])

        rev_table = doc.add_table(rows=1, cols=2)
        rev_table.style = 'Table Grid'
        rev_table.alignment = WD_TABLE_ALIGNMENT.LEFT
        _force_table_widths(rev_table, [2.25, 5.0])

        # Header
        for i, h in enumerate(['Binder Tab', 'Cards to Pull']):
            cell = rev_table.rows[0].cells[i]
            cell.text = h
            set_cell_shading(cell, '1B1F3B')
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.color.rgb = WHITE; r.bold = True
                    r.font.name = 'Arial'; r.font.size = Pt(8)

        for cat in ['People', 'Actions', 'Descriptions', 'Nouns', 'Prepositions', 'Social']:
            if cat not in by_cat:
                continue
            row = rev_table.add_row()
            row.cells[0].text = f'{FITZ_CATS.get(cat, cat)} ({cat})'
            row.cells[1].text = ', '.join(sorted(set(by_cat[cat])))
            for cell in row.cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.name = 'Arial'; r.font.size = Pt(8)

        _keep_table_together(rev_table)
        doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # ── HEART WORDS ──
    if heart_words:
        add_para(doc, f'Heart Words  ♥  ({len(hw_names)})', bold=True, size=11, color=TEAL, space_after=Pt(2))
        add_para(doc, 'Irregular words — cannot be fully decoded. Memorize by sight. Identify the "regular part" and the "heart part."',
                 size=8, color=GRAY, space_after=Pt(2))
        hw_text = '     '.join(f'♥ {w}' for w in hw_names)
        add_para(doc, hw_text, size=12, bold=True, space_after=Pt(4))

    # ══════════════════════════════════════════════════════════
    # PAGE 3: SESSION TRACKER (landscape partner handoff page)
    # ══════════════════════════════════════════════════════════
    build_tracker_page(doc, num, lesson_data=lesson)


# ══════════════════════════════════════════════════════════════
# Communication Partner Quick Reference (one page)
# ══════════════════════════════════════════════════════════════

def build_partner_quick_ref(doc):
    """One-page reference the partner can flip back to any time.

    Covers the 8 UFLI steps (what to do, what tool, what the student does),
    prompting reminders, wait time, and critical DO / DON'T rules.
    Dense but readable — this replaces having to find the Teacher Guide.
    """
    add_page_break(doc)
    add_heading(doc, 'Communication Partner Quick Reference', level=1)
    add_para(doc, 'Keep this page bookmarked. The routine is the same every lesson — only the words change.',
             size=9, color=GRAY, space_after=Pt(6))

    # ── THE 8 UFLI STEPS — compact table ──
    add_para(doc, 'The 8 Steps (same routine every lesson)', bold=True, size=11, color=NAVY, space_after=Pt(2))

    steps_data = [
        ('1. Phonemic\nAwareness', '~2 min', 'E-trans + symbol cards',
         '"I\'m going to say a word. Listen: [word]. Now select the sound you hear at the beginning."',
         'Selects from symbol cards or e-trans board'),
        ('2. Visual Drill', '~3 min', 'Letter cards or phoneme cards',
         '"Look at this letter. What sound does it make?" Show the letter, wait.',
         'Indicates the sound (partner voice, device, or pointing to phoneme card)'),
        ('3. Auditory Drill', '~5 min', 'Alternative pencil',
         '"I\'m going to say a sound. /[sound]/. Show me the letter that makes that sound."',
         'Selects letter on alternative pencil or letter card'),
        ('4. Blending Drill', '~5 min', 'E-trans + symbol cards',
         '"Let\'s blend this word. /[sounds]/... What word is that?" Point to each letter, say each sound.',
         'Selects the word from symbol cards on e-trans board'),
        ('5. New Concept', '~15 min', 'E-trans + alt. pencil',
         'Introduce today\'s phoneme/grapheme. Model the sound. "This letter says /[sound]/."',
         'Practices encoding (pencil) and decoding (e-trans) with new pattern'),
        ('6. Word Work', '~5 min', 'E-trans + letter tiles',
         '"Let\'s build words. Change [letter] to [letter]. What\'s the new word?"',
         'Manipulates letters (pencil/tiles) or selects word (e-trans)'),
        ('7. Heart Words', '~6 min', 'E-trans + alt. pencil',
         '"This is a heart word — we learn it by heart. The word is [word]. Show me [word]."',
         'Identifies heart word on e-trans. Traces/writes on alternative pencil'),
        ('8. Connected Text', '~15 min', 'E-trans + symbols + passage',
         '"Let\'s read together." Point to each word. Student decodes with support.',
         'Reads decodable passage with symbol support. Retells or answers comprehension'),
    ]

    st = doc.add_table(rows=len(steps_data) + 1, cols=5)
    st.style = 'Table Grid'
    st.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header
    for i, h in enumerate(['Step', 'Time', 'Tools', 'You Say', 'Student Does']):
        cell = st.rows[0].cells[i]
        cell.text = h
        set_cell_shading(cell, '1B1F3B')
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.color.rgb = WHITE; r.bold = True
                r.font.name = 'Arial'; r.font.size = Pt(7)

    for row_i, (step, time, tools, script, student) in enumerate(steps_data):
        row = st.rows[row_i + 1]
        shade = 'F4F6F8' if row_i % 2 == 0 else 'FFFFFF'
        texts = [step, time, tools, script, student]
        for ci, txt in enumerate(texts):
            row.cells[ci].text = txt
            set_cell_shading(row.cells[ci], shade)
            for p in row.cells[ci].paragraphs:
                for r in p.runs:
                    r.font.name = 'Arial'
                    r.font.size = Pt(7)
                    if ci == 0:
                        r.bold = True; r.font.color.rgb = NAVY

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # ── PROMPTING + WAIT TIME — two-column layout ──
    prompt_layout = doc.add_table(rows=1, cols=2)
    prompt_layout.alignment = WD_TABLE_ALIGNMENT.LEFT
    _remove_cell_borders(prompt_layout.rows[0].cells[0])
    _remove_cell_borders(prompt_layout.rows[0].cells[1])

    # LEFT: Prompting
    left = prompt_layout.rows[0].cells[0]
    lp = left.paragraphs[0]
    lr = lp.add_run('Prompting Reminders')
    lr.font.name = 'Arial'; lr.font.size = Pt(10); lr.bold = True; lr.font.color.rgb = NAVY

    prompts = [
        'Present → Wait → Prompt (never prompt before waiting)',
        'Wait time: 5–10 seconds of silence. Count in your head.',
        'If no response after wait: model the correct answer, then re-present.',
        'If incorrect: say "Let\'s try that again" — model, re-present.',
        'Fade prompts as student succeeds. Goal = independence.',
        'New skill: Full model first, then 3-sec delay, then 5-sec, then independent.',
        'Review skill: Start at the last successful prompt level.',
    ]
    for pr in prompts:
        pp = left.add_paragraph()
        prr = pp.add_run(f'• {pr}')
        prr.font.name = 'Arial'; prr.font.size = Pt(7)
        pp.paragraph_format.space_after = Pt(1)

    # RIGHT: DO / DON'T
    right = prompt_layout.rows[0].cells[1]
    rp = right.paragraphs[0]
    rr = rp.add_run('Do This / Not That')
    rr.font.name = 'Arial'; rr.font.size = Pt(10); rr.bold = True; rr.font.color.rgb = NAVY

    dos = [
        ('Wait in silence.', 'Fill the silence with talk.'),
        ('Say "Show me ___."', 'Say "Can you point to ___?"'),
        ('Accept any valid response mode.', 'Require one specific response.'),
        ('Model on the AAC system.', 'Only use your voice.'),
        ('Celebrate the attempt.', 'Only celebrate correct answers.'),
        ('Let the student struggle briefly.', 'Jump in at the first hesitation.'),
        ('Use the student\'s own AAC first.', 'Only use the provided symbols.'),
    ]

    do_table = doc.add_table(rows=len(dos) + 1, cols=2)
    do_table.style = 'Table Grid'
    do_table.alignment = WD_TABLE_ALIGNMENT.LEFT

    for i, h in enumerate(['DO', "DON'T"]):
        cell = do_table.rows[0].cells[i]
        cell.text = h
        set_cell_shading(cell, '1B1F3B')
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.color.rgb = WHITE; r.bold = True
                r.font.name = 'Arial'; r.font.size = Pt(7)

    for ri, (do, dont) in enumerate(dos):
        row = do_table.rows[ri + 1]
        row.cells[0].text = do
        row.cells[1].text = dont
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.name = 'Arial'; r.font.size = Pt(7)

    # Move the DO/DON'T table into the right cell
    right._tc.append(do_table._tbl)

    doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # ── AUDITORY LOOP reminder ──
    add_para(doc, 'The Auditory Loop (non-negotiable every step)',
             bold=True, size=9, color=NAVY, space_after=Pt(1))
    add_para(doc, 'Student selects a letter or word → Partner says the sound or word aloud → Student hears it → '
             'This builds the internal reading voice. If the student cannot hear the result of their selection, '
             'the phonics instruction is not working. Use partner voice, Talking Mirror, or SGD.',
             size=8, space_after=Pt(0))


# ══════════════════════════════════════════════════════════════
# TOC placeholder
# ══════════════════════════════════════════════════════════════

def add_toc(doc):
    """Add a Table of Contents field that Word will populate."""
    add_heading(doc, 'Table of Contents', level=1)

    # Add a TOC field — Word will update it when opened
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run._r.append(fldChar1)

    run2 = paragraph.add_run()
    instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> TOC \\o "1-2" \\h \\z \\u </w:instrText>')
    run2._r.append(instrText)

    run3 = paragraph.add_run()
    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
    run3._r.append(fldChar2)

    run4 = paragraph.add_run('Right-click → Update Field to populate this table of contents.')
    run4.font.color.rgb = GRAY
    run4.font.size = Pt(10)
    run4.font.name = 'Arial'

    run5 = paragraph.add_run()
    fldChar3 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run5._r.append(fldChar3)

    add_page_break(doc)


# ══════════════════════════════════════════════════════════════
# Main build
# ══════════════════════════════════════════════════════════════

def main():
    print('=' * 60)
    print('UFLI Foundations — Spiral-Bound Book Builder')
    print('9 Cent Copy Coil Binding | B&W Interior')
    print('=' * 60)

    # Load lesson data
    with open(LESSONS_JSON) as f:
        lessons = json.load(f)
    print(f'Loaded {len(lessons)} lesson configs')

    # Create document
    doc = Document()

    # ── Page setup: spiral binding margins ──
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.left_margin = Inches(0.75)   # Extra for coil punch
        section.right_margin = Inches(0.5)
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        # Mirror margins for double-sided
        section.gutter = Inches(0)
        # Set different odd/even headers if needed
        section.different_first_page_header_footer = False

    # ── Set default font ──
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    # Update heading styles
    for level in range(1, 4):
        hs = doc.styles[f'Heading {level}']
        hs.font.name = 'Arial'
        hs.font.color.rgb = NAVY

    # ══════════════════════════════════════════════════════════
    # COVER PAGE (placeholder — color cover printed separately)
    # ══════════════════════════════════════════════════════════
    cover_p = doc.add_paragraph()
    cover_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cover_p.paragraph_format.space_before = Pt(120)

    run = cover_p.add_run('COMMUNICATE BY DESIGN')
    run.font.size = Pt(24)
    run.font.color.rgb = TEAL
    run.bold = True
    run.font.name = 'Arial'

    doc.add_paragraph()
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run('UFLI Foundations')
    run.font.size = Pt(28)
    run.font.color.rgb = NAVY
    run.bold = True
    run.font.name = 'Arial'

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub_p.add_run('Complete AAC Access Guide')
    run.font.size = Pt(16)
    run.font.color.rgb = AMBER
    run.font.name = 'Arial'

    doc.add_paragraph()
    desc_p = doc.add_paragraph()
    desc_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = desc_p.add_run('Teacher Guide + Communication Partner Guide + Lessons 1–34')
    run.font.size = Pt(12)
    run.font.color.rgb = GRAY
    run.font.name = 'Arial'

    doc.add_paragraph()
    note_p = doc.add_paragraph()
    note_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = note_p.add_run('Where AT Meets Practice')
    run.font.size = Pt(11)
    run.font.color.rgb = GRAY
    run.font.name = 'Arial'
    run.italic = True

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════
    # TABLE OF CONTENTS
    # ══════════════════════════════════════════════════════════
    add_toc(doc)

    # ══════════════════════════════════════════════════════════
    # SECTION A + B: Teacher Guide + Communication Partner Guide
    # ══════════════════════════════════════════════════════════
    print('\n📖 Importing Teacher Guide + Communication Partner Guide...')
    import_teacher_guide(doc)

    # ══════════════════════════════════════════════════════════
    # PARTNER QUICK REFERENCE (one page, right before lessons)
    # ══════════════════════════════════════════════════════════
    print('\n📋 Building Communication Partner Quick Reference...')
    build_partner_quick_ref(doc)
    print('  ✅ Quick Reference page added')

    # ══════════════════════════════════════════════════════════
    # SECTION C: Lessons 1-34
    # ══════════════════════════════════════════════════════════
    print('\n📝 Building lesson pages...')

    # Section divider
    add_page_break(doc)
    divider_p = doc.add_paragraph()
    divider_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    divider_p.paragraph_format.space_before = Pt(200)
    run = divider_p.add_run('SECTION C')
    run.font.size = Pt(14)
    run.font.color.rgb = TEAL
    run.bold = True
    run.font.name = 'Arial'

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run('Lessons 1–34')
    run.font.size = Pt(24)
    run.font.color.rgb = NAVY
    run.bold = True
    run.font.name = 'Arial'

    note_p = doc.add_paragraph()
    note_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = note_p.add_run('Lessons 1–4: Letter introduction (no words — see Letter Card Library)\n'
                         'Lessons 5–34: Full lesson packets with symbol cards\n'
                         '★ = core word (high frequency)  |  ♥ = heart word (irregular)')
    run.font.size = Pt(10)
    run.font.color.rgb = GRAY
    run.font.name = 'Arial'

    for lesson in lessons:
        build_lesson_section(doc, lesson)
        print(f'  ✅ Lesson {lesson["number"]}')

    # ══════════════════════════════════════════════════════════
    # BACK MATTER
    # ══════════════════════════════════════════════════════════
    add_page_break(doc)
    add_heading(doc, 'Accessibility Statement', level=1)
    add_para(doc,
        'Communicate by Design is committed to creating accessible instructional '
        'materials. This document uses Arial font throughout, maintains consistent '
        'heading hierarchy for screen reader navigation, and avoids color-only '
        'information encoding. All tables include header rows.',
        size=10)

    doc.add_paragraph()
    add_heading(doc, 'About the Creator', level=1)
    add_para(doc,
        'Jill McCardel is a special educator and advocate, and the creator of '
        'Communicate by Design. With experience as both a SPED teacher and the parent '
        'of an AAC user, she builds the tools she wishes she had — evidence-based, '
        'practical resources that empower the whole team.',
        size=10)

    doc.add_paragraph()
    add_heading(doc, 'Attribution', level=1)
    add_para(doc,
        'UFLI Foundations is a product of the University of Florida Literacy Institute. '
        'This AAC access layer is an independent supplement — not affiliated with or '
        'endorsed by the University of Florida.',
        size=10)
    add_para(doc,
        'Fitzgerald Key classification system: Fitzgerald (1949), adapted by '
        "Goossens', Crain, & Elder (1992).",
        size=10)
    add_para(doc,
        '© Communicate by Design. All rights reserved.',
        size=10, bold=True)

    # ══════════════════════════════════════════════════════════
    # SAVE
    # ══════════════════════════════════════════════════════════
    output_path = os.path.join(OUTPUT_DIR,
        'UFLI_Complete_Spiral_Bound.docx')
    doc.save(output_path)

    # Count stats
    letter_lessons = [l for l in lessons if l['number'] < 5]
    word_lessons = [l for l in lessons if l['number'] >= 5]
    total_new = sum(len(l.get('newWords', [])) for l in word_lessons)
    total_review = sum(len(l.get('reviewWords', [])) for l in word_lessons)

    print(f'\n{"=" * 60}')
    print(f'✅ UFLI Complete Spiral-Bound Book')
    print(f'   Letter-only lessons (1–4): {len(letter_lessons)}')
    print(f'   Word lessons (5–34): {len(word_lessons)}')
    print(f'   Total lessons: {len(lessons)}')
    print(f'   New words across all lessons: {total_new}')
    print(f'   Review words across all lessons: {total_review}')
    print(f'   Sections: Cover + TOC + Teacher Guide + Comm Partner + 30 Lessons + Back Matter')
    print(f'   Interior: B&W (no symbol images)')
    print(f'   Binding: Left-edge coil (0.75" left margin)')
    print(f'   Output: {output_path}')
    print(f'{"=" * 60}')
    print(f'\nNext steps:')
    print(f'  1. Open in Word → right-click TOC → Update Field')
    print(f'  2. Create color cover in Canva (print separately)')
    print(f'  3. Save As → PDF for printer')
    print(f'  4. Send to 9 Cent Copy: Coil Bound Booklet, B&W interior, color cover')

    return output_path


if __name__ == '__main__':
    main()
