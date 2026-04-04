"""
CbD Shared Inserts Builder
==========================
Generates standalone insert pages that go into every CbD product:
  - Accessibility Statement
  - About the Creator
  - Terms of Use
  - Words to Add to the Device (UFLI only)

Run this to regenerate inserts after any approved content change.
Output: individual .docx files in this folder.

Usage:
  python3 build_shared_inserts.py           # builds all inserts
  python3 build_shared_inserts.py --insert accessibility
  python3 build_shared_inserts.py --insert about
  python3 build_shared_inserts.py --insert terms
  python3 build_shared_inserts.py --insert words_to_add

HARD RULE: Do NOT change the approved text in About, Terms, or Accessibility
without explicit instruction from Jill. These are locked.
"""

import argparse
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT_DIR = os.path.dirname(os.path.abspath(__file__))

# ── BRAND COLORS ──────────────────────────────────────────────────────────────
NAVY       = RGBColor(0x1B, 0x1F, 0x3B)
TEAL       = RGBColor(0x00, 0x6D, 0xA0)   # WCAG AA on white — docs only
AMBER      = RGBColor(0xFF, 0xB7, 0x03)
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
DARK_TEXT  = RGBColor(0x1A, 0x1A, 0x1A)

def set_heading_style(para, text, level=2, color=TEAL):
    """Set a heading paragraph with CbD teal color."""
    run = para.runs[0] if para.runs else para.add_run()
    run.text = text
    run.bold = True
    run.font.color.rgb = color
    run.font.size = Pt(13 if level == 2 else 11)

def add_heading(doc, text, level=2):
    p = doc.add_paragraph(style=f'Heading {level}')
    run = p.add_run(text)
    run.font.color.rgb = TEAL
    run.bold = True
    return p

def add_body(doc, text):
    p = doc.add_paragraph(style='Normal')
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = DARK_TEXT
    return p

def add_bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = DARK_TEXT
    return p

# ── ACCESSIBILITY STATEMENT ───────────────────────────────────────────────────
def build_accessibility(doc=None):
    if doc is None:
        doc = Document()
    add_heading(doc, 'Accessibility Statement', level=2)
    add_body(doc, (
        'This document was designed to meet WCAG 2.2 Level AA accessibility standards. '
        'It is compatible with screen readers, text-to-speech tools (including Microsoft '
        'Immersive Reader), and assistive technology used in special education settings.'
    ))
    add_body(doc, (
        'Using this document with assistive technology: In Microsoft Word, use '
        'View > Immersive Reader or Review > Read Aloud to hear passages read aloud. '
        'Reading passages are formatted for continuous text-to-speech use.'
    ))
    add_body(doc, (
        'Known limitation: If you convert this document to PDF, use '
        'File > Save As > PDF (not Print > PDF). The Save As method preserves document '
        'structure tags that screen readers need. Printing to PDF strips those tags.'
    ))
    add_body(doc, (
        'If you encounter an accessibility barrier in this product, contact Jill McCardel '
        'at communicatebydesign.substack.com. Accessibility is not an add-on — '
        'it is how Communicate by Design builds.'
    ))
    return doc

# ── ABOUT THE CREATOR ─────────────────────────────────────────────────────────
def build_about(doc=None):
    if doc is None:
        doc = Document()
    add_heading(doc, 'About the Creator', level=2)
    add_body(doc, (
        'Communicate by Design was created by Jill McCardel — '
        'a special educator, advocate, and mom to an AAC user.'
    ))
    add_body(doc, (
        'Everything in this resource is designed to build capacity across the whole team — '
        'teachers, OTs, PTs, BCBAs, RBTs, paraprofessionals, and families. '
        'AT and AAC aren\'t one person\'s job. Any professional who spends time '
        'with the student can learn to support communication.'
    ))
    add_body(doc, (
        'Communicate by Design is grounded in the capacity-building approach: '
        'evidence-based, practical, and ready to use. No fluff, no filler — '
        'just tools that work on Monday morning.'
    ))
    add_body(doc, (
        'Find more resources and writing at communicatebydesign.substack.com '
        'and on Instagram at @communicatebydesignaac.'
    ))
    return doc

# ── TERMS OF USE ──────────────────────────────────────────────────────────────
def build_terms(doc=None):
    if doc is None:
        doc = Document()
    add_heading(doc, 'Terms of Use', level=2)
    add_body(doc, 'Thank you for purchasing from Communicate by Design!')
    add_body(doc, 'By purchasing this resource, you agree to the following terms of use.')
    add_heading(doc, 'What You CAN Do', level=3)
    for item in [
        'Use this resource in your own classroom, therapy room, or home setting with your own students or child.',
        'Print or photocopy pages for use with your students or child.',
        'Share the resource with a student\'s IEP team for direct use with that student '
        '(e.g., sending a copy to a parent, OT, PT, BCBA, RBT, or paraprofessional working with the same learner).',
        'Save a digital backup copy for your personal use.',
    ]:
        add_bullet(doc, item)
    add_heading(doc, 'What You CANNOT Do', level=3)
    for item in [
        'Share this resource with other teachers, therapists, or colleagues for use in their own classrooms. '
        'Each teacher or professional needs their own license.',
        'Post this resource — or any part of it — on any website, shared drive, or online platform '
        '(including school or district servers, Google Drive shared folders, or learning management systems '
        'accessible by others).',
        'Sell, redistribute, or claim this resource as your own.',
        'Edit, modify, or remove the copyright information from any part of this resource.',
    ]:
        add_bullet(doc, item)
    add_body(doc, (
        'School-Wide or District Licenses: Need this resource for your whole team? '
        'I offer discounted multi-license options. Contact me through my TPT store for pricing — '
        'I\'m happy to work with schools and districts.'
    ))
    add_body(doc, (
        '© Communicate by Design. All rights reserved. Purchase of this resource grants a '
        'single-user license only. This resource may not be reproduced, distributed, or '
        'displayed publicly without written permission.'
    ))
    return doc

# ── WORDS TO ADD TO THE DEVICE (UFLI) ────────────────────────────────────────
def build_words_to_add(doc=None):
    if doc is None:
        doc = Document()
    add_heading(doc, 'Words to Add to the Device', level=1)
    add_body(doc, (
        'Vocabulary exploration is ongoing — not a one-time setup task. '
        'Words from each lesson belong in the student\'s communication system because '
        'the student is learning to read those words. The goal is not device compliance. '
        'The goal is that the student encounters these words in reading AND in communication '
        '— and that every adult on the team is modeling those words continuously.'
    ))
    # Student-led callout
    p = doc.add_paragraph()
    run = p.add_run(
        'STUDENT-LED MEANS: The student decides when and how to use a new word. '
        'Your job is to make it available, model it frequently, and celebrate any use — '
        'spontaneous, prompted, or exploratory. A student who selects a word \'wrong\' '
        'is still using language. That is always the goal.'
    )
    run.bold = True
    run.font.size = Pt(10)

    add_heading(doc, 'Before the Lesson', level=2)
    add_body(doc, (
        'Review the per-lesson packet word list. For each new word, check: '
        'Is it already in the student\'s system? If yes — great, no action needed. '
        'If no — add it, or flag it for the SLP/AAC team. Core words should already '
        'be on the system. Fringe words are lesson-specific and need to be added before instruction.'
    ))

    add_heading(doc, 'During Every Lesson — Model Continuously', level=2)
    add_body(doc, (
        'Every time a target word appears in the lesson, model it on the student\'s system. '
        'Do not wait for the student to use it first. Do not prompt the student to use it. '
        'Modeling is not prompting — you are showing, not demanding. '
        'Say the word, find it on the device or e-trans board, and touch it. Then continue.'
    ))

    add_heading(doc, 'Who Adds Words?', level=2)
    add_body(doc, (
        'Anyone on the team — the student, the communication partner, the special educator, '
        'the SLP, a family member. This is not an SLP-only task. '
        'If a word is in a lesson the student is doing today, it belongs on the system today. '
        'The SLP is a collaborator and resource — not the gatekeeper for vocabulary access.'
    ))

    add_heading(doc, 'Beyond the Lesson — Keep Exploring', level=2)
    add_body(doc, (
        'Lesson words are a starting point. If the student starts using a lesson word '
        'in a new context — outside of phonics time, with a different partner, to comment '
        'rather than respond — that is language growth. Note it. Celebrate it. '
        'Use the Session Data Tracker (in each lesson packet) to record spontaneous use '
        'and generalization across partners. These observations are the evidence that '
        'language is growing, not just compliance data.'
    ))
    return doc

# ── MAIN ──────────────────────────────────────────────────────────────────────
def build_all():
    inserts = {
        'CbD_Insert_Accessibility.docx':   build_accessibility,
        'CbD_Insert_About.docx':           build_about,
        'CbD_Insert_Terms.docx':           build_terms,
        'CbD_Insert_Words_to_Add.docx':    build_words_to_add,
    }
    for filename, builder in inserts.items():
        doc = builder()
        path = os.path.join(OUT_DIR, filename)
        doc.save(path)
        print(f'Saved: {filename}')
    print('\nAll inserts built.')

if __name__ == '__main__':
    parser = argparse.ArgumentParser()
    parser.add_argument('--insert', choices=['accessibility','about','terms','words_to_add'], default=None)
    args = parser.parse_args()

    if args.insert == 'accessibility':
        d = build_accessibility(); d.save(os.path.join(OUT_DIR,'CbD_Insert_Accessibility.docx')); print('Saved.')
    elif args.insert == 'about':
        d = build_about(); d.save(os.path.join(OUT_DIR,'CbD_Insert_About.docx')); print('Saved.')
    elif args.insert == 'terms':
        d = build_terms(); d.save(os.path.join(OUT_DIR,'CbD_Insert_Terms.docx')); print('Saved.')
    elif args.insert == 'words_to_add':
        d = build_words_to_add(); d.save(os.path.join(OUT_DIR,'CbD_Insert_Words_to_Add.docx')); print('Saved.')
    else:
        build_all()
